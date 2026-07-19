#!/usr/bin/env python3
"""
ATS Resume Score Optimizer — Hybrid (rule-based + Gemini AI)

GENERAL MODE (no job description):
    python ats_optimizer.py --resume resume.txt

JD MATCH MODE:
    python ats_optimizer.py --resume resume.docx --jd jd.txt

AI LAYER (recommended for the most accurate score):
    Get a free key at https://aistudio.google.com/apikey, then:

        export GEMINI_API_KEY=AIza...
        python ats_optimizer.py --resume resume.docx --jd jd.txt

    Without a key the tool still runs on rule-based analysis alone.
    Disable AI even when a key is set: --no-ai
    Switch model: --model gemini-3.5-flash  (default: gemini-2.5-flash)

FLAGS:
    --resume PATH      .txt/.md/.pdf/.docx, or raw text with --resume-text
    --jd PATH          optional JD file/text; omit for general mode
    --resume-text      treat --resume value as raw text
    --jd-text          treat --jd value as raw text
    --no-ai            skip the AI layer even if a key is available
    --model NAME       Gemini model string (default: gemini-2.5-flash)
    --api-key KEY      Gemini API key (else reads GEMINI_API_KEY env var)
    --json             machine-readable JSON output

OPTIONAL DEPS:
    pip install python-docx PyPDF2 scikit-learn --break-system-packages
    - python-docx : read .docx + detect ATS-breaking structures
    - PyPDF2      : read .pdf
    - scikit-learn: TF-IDF similarity bonus in JD match mode
"""

import re
import os
import sys
import json
import argparse
import urllib.request
import urllib.error
from collections import Counter
import ssl

# Fix for macOS Python SSL Certificate Verify Failed error
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

try:
    import docx  # python-docx
except ImportError:
    docx = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False


# Resume analysis is quality-critical, so use the stronger FLASH model. Its free
# daily bucket is reserved for quality work (cover letters, resume tips, ranking,
# ATS) now that the high-volume email classifier runs on flash-lite instead.
DEFAULT_MODEL = "gemini-2.5-flash"
GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta/models"


# =========================================================================
# Rule-based constants
# =========================================================================

STOPWORDS = set("""
a about above after again against all am an and any are aren't as at be
because been before being below between both but by can't cannot could
couldn't did didn't do does doesn't doing don't down during each few for
from further had hadn't has hasn't have haven't having he he'd he'll he's
her here here's hers herself him himself his how how's i i'd i'll i'm i've
if in into is isn't it it's its itself let's me more most mustn't my
myself no nor not of off on once only or other ought our ours ourselves
out over own same shan't she she'd she'll she's should shouldn't so some
such than that that's the their theirs them themselves then there there's
these they they'd they'll they're they've this those through to too under
until up very was wasn't we we'd we'll we're we've were weren't what
what's when when's where where's which while who who's whom why why's
with won't would wouldn't you you'd you'll you're you've your yours
yourself yourselves etc per via using use used within across including
ability able strong excellent good great new also will must can may
""".split())

ACTION_VERBS = set("""
achieved accelerated administered advised analyzed architected automated
boosted built calculated centralized chaired coached collaborated
conceptualized conducted coordinated created cultivated decreased
delivered designed developed devised directed drove earned engineered
enhanced established evaluated executed expanded facilitated forecasted
formulated founded generated grew guided headed identified implemented
improved increased influenced initiated innovated instituted integrated
introduced launched led leveraged liaised managed maximized mentored
migrated minimized modernized motivated negotiated operated optimized
orchestrated organized overhauled oversaw partnered performed pioneered
planned prepared presented prioritized produced programmed projected
promoted proposed published recruited reduced refined reorganized
researched resolved restructured revamped reviewed scaled secured shaped
simplified slashed spearheaded standardized steered streamlined
strengthened structured supervised supported surpassed trained
transformed translated unified upgraded validated won wrote
""".split())

WEAK_PHRASES = [
    "references available upon request",
    "responsible for",
    "duties included",
    "duties include",
    "team player",
    "hard worker",
    "go-getter",
    "think outside the box",
    "synergy",
    "self-starter",
    "detail oriented",
]

PERSONAL_PRONOUNS = {"i", "me", "my", "mine", "myself", "we", "our", "ours"}

SECTION_PATTERNS = {
    "summary":        [r"\bsummary\b", r"\bobjective\b", r"\bprofile\b", r"\babout me\b"],
    "experience":     [r"experience", r"employment", r"work history", r"professional background"],
    "education":      [r"education", r"academic", r"qualifications"],
    "skills":         [r"skills", r"competenc", r"technical proficienc", r"core strengths"],
    "projects":       [r"\bprojects\b", r"portfolio"],
    "certifications": [r"certificat", r"licens"],
}

DATE_PATTERNS = [
    r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\b",
    r"\b\d{1,2}/\d{4}\b",
    r"\b\d{4}\s*[-\u2013\u2014]\s*\d{4}\b",
    r"\b\d{4}\s*[-\u2013\u2014]\s*(present|current)\b",
    r"\b\d{1,2}/\d{4}\s*[-\u2013\u2014]\s*(\d{1,2}/\d{4}|present|current)\b",
]

NORMALIZATIONS = [
    (r"asp\.net", "aspnet"),
    (r"\.net\b",  "dotnet"),
    (r"c\+\+",    "cplusplus"),
    (r"c#",       "csharp"),
    (r"f#",       "fsharp"),
    (r"node\.js", "nodejs"),
    (r"vue\.js",  "vuejs"),
    (r"react\.js","reactjs"),
    (r"ci/cd",    "cicd"),
]

SKILL_HINTS = set("""
python java javascript typescript cplusplus csharp golang rust ruby php
swift kotlin scala sql nosql html css react angular vue node django flask
spring fastapi express rails laravel aws azure gcp docker kubernetes
terraform ansible jenkins git github gitlab cicd linux bash agile scrum
kanban jira confluence graphql microservices tensorflow pytorch nlp etl
tableau powerbi excel salesforce hubspot seo sem figma sketch photoshop
illustrator selenium pytest junit sap oracle workday dotnet aspnet nodejs
fsharp gaap b2b b2c crm erp ux ui wireframing prototyping cybersecurity
devops
""".split())


# =========================================================================
# File reading + DOCX structural ATS check
# =========================================================================

def read_docx(path):
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(p for p in parts if p)


def check_docx_ats_issues(path):
    """Return structural red flags that often break ATS parsing."""
    if docx is None:
        return []
    try:
        document = docx.Document(path)
    except Exception:
        return []

    issues = []
    if document.tables:
        issues.append(
            f"Contains {len(document.tables)} table(s) — many ATS parsers misread or "
            "skip table content; use plain text with consistent layout instead."
        )
    if document.inline_shapes:
        issues.append(
            f"Contains {len(document.inline_shapes)} image(s)/icon(s) — ATS cannot read "
            "text in images, and photos can trigger bias filters in some systems."
        )
    header_footer_text = False
    for section in document.sections:
        if any(p.text.strip() for p in section.header.paragraphs):
            header_footer_text = True
        if any(p.text.strip() for p in section.footer.paragraphs):
            header_footer_text = True
    if header_footer_text:
        issues.append(
            "Contact info or other content found in a header/footer — many ATS "
            "parsers ignore headers and footers entirely."
        )
    try:
        if "txbxContent" in document.element.xml:
            issues.append(
                "Contains text box(es) — ATS parsers typically cannot extract "
                "text from text boxes at all."
            )
    except Exception:
        pass
    return issues


def read_text_input(value, is_path):
    if not is_path:
        return value

    ext = os.path.splitext(value)[1].lower()

    if ext in (".txt", ".md"):
        with open(value, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".pdf":
        if PyPDF2 is None:
            sys.exit("PyPDF2 is required for PDF files:  pip install PyPDF2 --break-system-packages")
        text = []
        with open(value, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)

    if ext == ".docx":
        if docx is None:
            sys.exit("python-docx is required for DOCX files:  pip install python-docx --break-system-packages")
        return read_docx(value)

    sys.exit(f"Unsupported file type: {ext}")


# =========================================================================
# Text utilities
# =========================================================================

WORD_RE = re.compile(r"[a-zA-Z][a-zA-Z0-9]*")


def normalize(text):
    text = text.lower()
    for pattern, repl in NORMALIZATIONS:
        text = re.sub(pattern, repl, text)
    return text


def tokenize(text):
    return WORD_RE.findall(normalize(text))


def content_words(text):
    return [t for t in tokenize(text) if t not in STOPWORDS and len(t) > 1]


def ngrams(tokens, n):
    return [" ".join(tokens[i:i + n]) for i in range(len(tokens) - n + 1)]


def get_bullets(text):
    bullets = []
    for line in text.splitlines():
        line = line.strip()
        if re.match(r"^[\u2022\-*\u25cf\u25aa\u25e6\u2023]\s+", line) or re.match(r"^\d+[.)]\s+", line):
            bullets.append(re.sub(r"^[\u2022\-*\u25cf\u25aa\u25e6\u2023\d.)\s]+", "", line))
    return bullets


def _stem(word):
    for suf in ("ing", "ed", "es", "s", "d"):
        if word.endswith(suf) and len(word) - len(suf) >= 3:
            return word[:-len(suf)]
    return word


ACTION_VERB_STEMS = {_stem(v) for v in ACTION_VERBS} | ACTION_VERBS


def starts_with_action_verb(bullet):
    words = bullet.split()
    if not words:
        return False
    first = re.sub(r"[^a-zA-Z]", "", words[0].lower())
    if not first:
        return False
    return first in ACTION_VERBS or _stem(first) in ACTION_VERB_STEMS


# =========================================================================
# Rule-based: general ATS score (no JD)
# =========================================================================

def score_general(resume_text):
    score = 0
    max_score = 100
    checks = []
    suggestions = []

    text_lower = resume_text.lower()
    bullets     = get_bullets(resume_text)
    words       = tokenize(resume_text)
    word_count  = len(words)

    # 1. Contact info (10 pts)
    has_email = bool(re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", resume_text))
    has_phone = bool(re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}", resume_text))
    pts = (5 if has_email else 0) + (5 if has_phone else 0)
    if not has_email: suggestions.append("Add a professional email address near the top.")
    if not has_phone: suggestions.append("Add a phone number near the top.")
    score += pts; checks.append(("Contact information", pts, 10))

    # 2. Core sections (15 pts)
    sections_found = [s for s, patterns in SECTION_PATTERNS.items()
                       if any(re.search(p, text_lower) for p in patterns)]
    core_sections = {"experience", "education", "skills"}
    found_core = core_sections & set(sections_found)
    pts = round(15 * (len(found_core) / len(core_sections)))
    score += pts; checks.append(("Standard sections (Experience/Education/Skills)", pts, 15))
    missing_core = core_sections - found_core
    if missing_core:
        suggestions.append(f"Add clearly labeled section(s) for: {', '.join(sorted(missing_core)).title()}.")

    # 3. Summary/Objective (5 pts)
    pts = 5 if "summary" in sections_found else 0
    score += pts; checks.append(("Summary/Objective section", pts, 5))
    if pts == 0: suggestions.append("Consider adding a 2-3 line professional summary at the top.")

    # 4. Bullet usage (10 pts)
    pts = 10 if len(bullets) >= 4 else round(10 * len(bullets) / 4)
    score += pts; checks.append(("Bullet point usage", pts, 10))
    if len(bullets) < 4:
        suggestions.append("Use bullet points for responsibilities and achievements (aim for 3-6 per role).")

    # 5. Action verbs (15 pts)
    if bullets:
        ratio = sum(starts_with_action_verb(b) for b in bullets) / len(bullets)
        pts = round(15 * ratio)
    else:
        pts = 0
    score += pts; checks.append(("Bullets start with action verbs", pts, 15))
    if bullets and pts < 12:
        suggestions.append('Start bullets with strong action verbs ("Led", "Built", "Increased") not "Responsible for".')

    # 6. Quantified achievements (15 pts)
    if bullets:
        ratio = sum(1 for b in bullets if re.search(r"\d", b)) / len(bullets)
        pts = round(15 * ratio)
    else:
        pts = 0
    score += pts; checks.append(("Quantified achievements (numbers/%/$)", pts, 15))
    if bullets and pts < 12:
        suggestions.append('Quantify achievements (e.g. "Increased revenue by 23%", "Managed $2M budget").')

    # 7. Length (10 pts)
    if 350 <= word_count <= 900:   pts = 10
    elif 250 <= word_count < 350:  pts = 7
    elif 900 < word_count <= 1100: pts = 7
    elif word_count < 250:         pts = 3
    else:                          pts = 5
    score += pts; checks.append((f"Length ({word_count} words)", pts, 10))
    if word_count < 250:  suggestions.append("Resume is too short — add detail on responsibilities and achievements.")
    if word_count > 1100: suggestions.append("Resume is too long — aim for 400-800 words (1-2 pages).")

    # 8. Clichés (5 pts)
    found_weak = [p for p in WEAK_PHRASES if p in text_lower]
    pts = max(0, 5 - len(found_weak))
    score += pts; checks.append(("Avoids clichés/weak phrases", pts, 5))
    if found_weak: suggestions.append(f"Remove or rewrite: {', '.join(found_weak)}.")

    # 9. Date consistency (5 pts)
    found_formats = {i for i, p in enumerate(DATE_PATTERNS) if re.search(p, text_lower)}
    pts = 5 if len(found_formats) <= 1 else (3 if len(found_formats) == 2 else 1)
    score += pts; checks.append(("Consistent date formatting", pts, 5))
    if len(found_formats) > 1:
        suggestions.append('Use one date format throughout (e.g. "Jan 2020 – Present").')

    # 10. No pronouns (5 pts)
    pronoun_count = sum(1 for w in words if w in PERSONAL_PRONOUNS)
    pts = max(0, 5 - pronoun_count)
    score += pts; checks.append(("Avoids first-person pronouns", pts, 5))
    if pronoun_count > 0:
        suggestions.append('Remove personal pronouns ("I", "my", "we") — resumes use implied first person.')

    # 11. Clean characters (5 pts)
    weird_chars = [c for c in re.findall(r"[^\x00-\x7F]", resume_text)
                   if c not in "\u2022\u25cf\u25aa\u25e6\u2023\u2013\u2014\u2019\u2018\u201c\u201d\u2026"]
    pts = 5 if not weird_chars else max(0, 5 - len(set(weird_chars)))
    score += pts; checks.append(("No unusual characters/symbols", pts, 5))
    if weird_chars:
        suggestions.append("Remove unusual symbols that some ATS parsers may misread.")

    return {
        "score": score,
        "max_score": max_score,
        "checks": [{"name": n, "points": p, "max": m} for n, p, m in checks],
        "sections_found": sections_found,
        "suggestions": suggestions,
    }


# =========================================================================
# Rule-based: JD match score
# =========================================================================

def extract_keywords(text, top_n=40):
    words   = content_words(text)
    bigrms  = ngrams(words, 2)
    scored  = Counter()
    for w, c in Counter(words).items():
        scored[w] = c + (3 if w in SKILL_HINTS else 0)
    for b, c in Counter(bigrms).items():
        parts = b.split()
        if c >= 2 or any(p in SKILL_HINTS for p in parts):
            scored[b] = c * 1.5 + (3 if any(p in SKILL_HINTS for p in parts) else 0)
    return [kw for kw, _ in scored.most_common(top_n)]


def score_jd_match(resume_text, jd_text):
    resume_lower = resume_text.lower()
    jd_keywords  = extract_keywords(jd_text, top_n=40)

    matched, missing = [], []
    for kw in jd_keywords:
        if kw in resume_lower or all(p in resume_lower for p in kw.split()):
            matched.append(kw)
        else:
            missing.append(kw)

    keyword_coverage = len(matched) / len(jd_keywords) if jd_keywords else 1.0

    similarity = None
    if SKLEARN_AVAILABLE:
        try:
            tfidf      = TfidfVectorizer(stop_words="english").fit_transform([resume_text, jd_text])
            similarity = float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0])
        except ValueError:
            similarity = None

    if similarity is not None:
        final_score = round((0.65 * keyword_coverage + 0.35 * similarity) * 100)
    else:
        final_score = round(keyword_coverage * 100)

    jd_skills     = set(content_words(jd_text)) & SKILL_HINTS
    missing_skills = sorted(jd_skills - set(content_words(resume_text)))

    suggestions = []
    if missing[:10]:
        suggestions.append("Weave in these missing JD keywords: " + ", ".join(missing[:10]) + ".")
    if missing_skills:
        suggestions.append("These tools/skills appear in the JD but not your resume (add if relevant): "
                           + ", ".join(missing_skills[:10]) + ".")
    if keyword_coverage < 0.5:
        suggestions.append("Keyword overlap is low — mirror the JD's exact language in your skills/experience sections.")

    return {
        "score":                final_score,
        "keyword_coverage_pct": round(keyword_coverage * 100),
        "similarity_pct":       round(similarity * 100) if similarity is not None else None,
        "matched_keywords":     matched,
        "missing_keywords":     missing,
        "missing_skills":       missing_skills,
        "suggestions":          suggestions,
    }


# =========================================================================
# Gemini AI layer
# =========================================================================

AI_INSTRUCTION = (
    "You are an expert ATS analyst and resume coach. "
    "Respond with ONLY a single valid JSON object exactly matching the "
    "requested schema — no markdown fences, no preamble, no extra text."
)

GENERAL_SCHEMA = """{
  "content_quality_score": <int 0-100, based on impact/clarity/achievement-focus/readability>,
  "inferred_role": "<best-fit job title or field for this resume>",
  "suggested_ats_keywords": ["<10-15 ATS-relevant keywords/skills for that role>"],
  "bullet_improvements": [
    {"original": "<weak bullet verbatim>", "improved": "<rewritten with stronger verb and/or quantification>", "reason": "<why the change helps>"}
  ],
  "strengths":  ["<2-4 specific strengths>"],
  "weaknesses": ["<2-4 specific weaknesses>"],
  "summary": "<2-3 sentence overall assessment>"
}"""

JD_SCHEMA = """{
  "semantic_match_score": <int 0-100, how well the resume's real experience aligns with the JD's core requirements, counting synonyms/related skills as partial matches>,
  "matched_requirements": ["<JD requirement> -> <how resume addresses it>"],
  "missing_requirements": ["<important JD requirements absent from the resume>"],
  "synonym_matches":       ["<JD term> -> <equivalent found in resume>"],
  "bullet_improvements": [
    {"original": "<bullet verbatim>", "improved": "<rewritten to mirror JD language and show impact>", "reason": "<why>"}
  ],
  "tailoring_suggestions": ["<specific actions to improve this resume for this exact JD>"],
  "summary": "<2-3 sentence overall fit assessment>"
}"""


def call_gemini(prompt_text, model, api_key, max_tokens=16384):
    url = f"{GEMINI_API_BASE}/{model}:generateContent"
    body = json.dumps({
        "contents": [{"role": "user", "parts": [{"text": prompt_text}]}],
        "systemInstruction": {"parts": [{"text": AI_INSTRUCTION}]},
        "generationConfig": {
            "maxOutputTokens": max_tokens,
            "temperature": 0.2,
            "responseMimeType": "application/json",
            # gemini-2.5-flash is a "thinking" model and thinking tokens count
            # against maxOutputTokens. Previously thinking ate the 4096 budget and
            # truncated the JSON (-> "Unterminated string"). Keep reasoning but
            # CAP it, inside a larger ceiling, so thinking + the full JSON both fit.
            "thinkingConfig": {"thinkingBudget": 2048},
        },
    }).encode("utf-8")

    req = urllib.request.Request(
        url,
        data=body,
        headers={
            "x-goog-api-key": api_key,
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Gemini API error {e.code}: {detail[:400]}") from e
    except urllib.error.URLError as e:
        raise RuntimeError(f"Network error calling Gemini API: {e}") from e

    try:
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError) as e:
        raise RuntimeError(f"Unexpected Gemini response shape: {json.dumps(data)[:300]}") from e


def parse_ai_json(raw_text):
    text = raw_text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\s*", "", text)
        text = re.sub(r"\s*```$", "", text.strip())
    return json.loads(text.strip())


def run_ai_general(resume_text, model, api_key):
    prompt = (
        "Analyze the following resume for ATS-friendliness and content quality. "
        f"Return ONLY JSON matching this schema:\n\n{GENERAL_SCHEMA}\n\n"
        f"RESUME:\n{resume_text}"
    )
    return parse_ai_json(call_gemini(prompt, model, api_key))


def run_ai_jd(resume_text, jd_text, model, api_key):
    prompt = (
        "Compare the resume against the job description for ATS and recruiter alignment. "
        f"Return ONLY JSON matching this schema:\n\n{JD_SCHEMA}\n\n"
        f"JOB DESCRIPTION:\n{jd_text}\n\nRESUME:\n{resume_text}"
    )
    return parse_ai_json(call_gemini(prompt, model, api_key))


def _safe_score(value, default):
    try:
        return max(0, min(100, int(round(float(value)))))
    except (TypeError, ValueError):
        return default


def _safe_list(value):
    return value if isinstance(value, list) else []


# =========================================================================
# Hybrid scoring
# =========================================================================

def run_general(resume_text, docx_issues, model, api_key, use_ai):
    rules  = score_general(resume_text)
    result = {"mode": "general", "rule_based": rules, "ai": None,
              "ai_error": None, "docx_issues": docx_issues}

    final = rules["score"]
    if use_ai:
        try:
            ai = run_ai_general(resume_text, model, api_key)
            result["ai"] = ai
            final = round(0.45 * rules["score"] + 0.55 * _safe_score(ai.get("content_quality_score"), rules["score"]))
        except Exception as e:
            result["ai_error"] = str(e)

    final = max(0, final - min(15, 5 * len(docx_issues)))
    result["final_score"] = final
    return result


def run_jd(resume_text, jd_text, docx_issues, model, api_key, use_ai):
    rules  = score_jd_match(resume_text, jd_text)
    result = {"mode": "jd_match", "rule_based": rules, "ai": None,
              "ai_error": None, "docx_issues": docx_issues}

    final = rules["score"]
    if use_ai:
        try:
            ai = run_ai_jd(resume_text, jd_text, model, api_key)
            result["ai"] = ai
            final = round(0.3 * rules["score"] + 0.7 * _safe_score(ai.get("semantic_match_score"), rules["score"]))
        except Exception as e:
            result["ai_error"] = str(e)

    final = max(0, final - min(15, 5 * len(docx_issues)))
    result["final_score"] = final
    return result


# =========================================================================
# Report printing
# =========================================================================

def _print_docx_issues(issues):
    if issues:
        print("\nATS PARSEABILITY WARNINGS (.docx structure):")
        for i in issues:
            print(f"  ! {i}")


def print_general_report(r):
    rb, ai = r["rule_based"], r["ai"]
    print("=" * 64)
    print("ATS RESUME ANALYSIS — GENERAL MODE")
    print("=" * 64)
    print(f"\nFINAL SCORE: {r['final_score']} / 100")
    print(f"  Rule-based formatting/structure score : {rb['score']}/100")
    if ai:
        print(f"  Gemini content-quality score          : {_safe_score(ai.get('content_quality_score'), rb['score'])}/100")
    elif r["ai_error"]:
        print(f"  AI analysis unavailable: {r['ai_error']}")

    print("\nFormatting / structure breakdown:")
    for c in rb["checks"]:
        bar = "█" * c["points"] + "░" * (c["max"] - c["points"])
        print(f"  {c['name']:<48} {bar}  {c['points']}/{c['max']}")

    print("\nSections detected:", ", ".join(rb["sections_found"]) or "none")
    _print_docx_issues(r["docx_issues"])

    if ai:
        print(f"\nInferred target role  : {ai.get('inferred_role', 'n/a')}")
        print("\nStrengths:")
        for s in _safe_list(ai.get("strengths")):
            print(f"  ✓ {s}")
        print("\nWeaknesses:")
        for w in _safe_list(ai.get("weaknesses")):
            print(f"  ✗ {w}")
        kws = _safe_list(ai.get("suggested_ats_keywords"))
        if kws:
            print("\nSuggested ATS keywords for this role:")
            print("  " + ", ".join(kws))
        bullets = _safe_list(ai.get("bullet_improvements"))
        if bullets:
            print("\nBullet rewrite suggestions:")
            for b in bullets:
                print(f"\n  Original : {b.get('original')}")
                print(f"  Improved : {b.get('improved')}")
                print(f"  Why      : {b.get('reason')}")
        if ai.get("summary"):
            print(f"\nAssessment: {ai['summary']}")

    if rb["suggestions"]:
        print("\nRule-based suggestions:")
        for s in rb["suggestions"]:
            print(f"  → {s}")


def print_jd_report(r):
    rb, ai = r["rule_based"], r["ai"]
    print("=" * 64)
    print("ATS RESUME ANALYSIS — JD MATCH MODE")
    print("=" * 64)
    print(f"\nFINAL MATCH SCORE: {r['final_score']} / 100")
    sim_str = f", TF-IDF similarity {rb['similarity_pct']}%" if rb["similarity_pct"] is not None else ""
    print(f"  Rule-based score   : {rb['score']}/100  "
          f"(keyword coverage {rb['keyword_coverage_pct']}%{sim_str})")
    if ai:
        print(f"  Gemini semantic match score : {_safe_score(ai.get('semantic_match_score'), rb['score'])}/100")
    elif r["ai_error"]:
        print(f"  AI analysis unavailable: {r['ai_error']}")

    _print_docx_issues(r["docx_issues"])

    if ai:
        if ai.get("summary"):
            print(f"\nAssessment: {ai['summary']}")
        matched = _safe_list(ai.get("matched_requirements"))
        if matched:
            print("\nMatched requirements:")
            for m in matched:
                print(f"  ✓ {m}")
        missing = _safe_list(ai.get("missing_requirements"))
        if missing:
            print("\nMissing requirements:")
            for m in missing:
                print(f"  ✗ {m}")
        synonyms = _safe_list(ai.get("synonym_matches"))
        if synonyms:
            print("\nSynonym matches (AI detected these equivalences):")
            for s in synonyms:
                print(f"  ~ {s}")
        bullets = _safe_list(ai.get("bullet_improvements"))
        if bullets:
            print("\nBullet rewrite suggestions (tailored to this JD):")
            for b in bullets:
                print(f"\n  Original : {b.get('original')}")
                print(f"  Improved : {b.get('improved')}")
                print(f"  Why      : {b.get('reason')}")
        tips = _safe_list(ai.get("tailoring_suggestions"))
        if tips:
            print("\nTailoring suggestions:")
            for s in tips:
                print(f"  → {s}")

    print(f"\nRule-based keyword match ({len(rb['matched_keywords'])}/{len(rb['matched_keywords'])+len(rb['missing_keywords'])}):")
    print("  Matched : " + (", ".join(rb["matched_keywords"][:20]) or "none"))
    print("  Missing : " + (", ".join(rb["missing_keywords"][:20]) or "none"))

    if rb["suggestions"]:
        print("\nRule-based suggestions:")
        for s in rb["suggestions"]:
            print(f"  → {s}")


# =========================================================================
# CLI
# =========================================================================

def main():
    try:
        import signal
        signal.signal(signal.SIGPIPE, signal.SIG_DFL)
    except (AttributeError, ValueError):
        pass

    parser = argparse.ArgumentParser(
        description="Hybrid (rules + Gemini AI) ATS Resume Score Optimizer",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--resume",      required=True, help="Resume file (.txt/.pdf/.docx) or raw text")
    parser.add_argument("--jd",          default=None,  help="Job description file/text (omit for general mode)")
    parser.add_argument("--resume-text", action="store_true", help="Treat --resume as raw text")
    parser.add_argument("--jd-text",     action="store_true", help="Treat --jd as raw text")
    parser.add_argument("--no-ai",       action="store_true", help="Skip Gemini AI layer")
    parser.add_argument("--model",       default=DEFAULT_MODEL, help=f"Gemini model (default: {DEFAULT_MODEL})")
    parser.add_argument("--api-key",     default=None,  help="Gemini API key (else reads GEMINI_API_KEY)")
    parser.add_argument("--json",        action="store_true", help="Output JSON")
    args = parser.parse_args()

    resume_text = read_text_input(args.resume, is_path=not args.resume_text)
    if not resume_text.strip():
        sys.exit("Resume text is empty.")

    docx_issues = []
    if not args.resume_text and os.path.splitext(args.resume)[1].lower() == ".docx":
        docx_issues = check_docx_ats_issues(args.resume)

    api_key = args.api_key or os.environ.get("GEMINI_API_KEY")
    use_ai  = bool(api_key) and not args.no_ai

    if not use_ai and not args.no_ai and not args.json:
        print("ℹ  AI layer disabled — set GEMINI_API_KEY (free at https://aistudio.google.com/apikey)")
        print("   to unlock semantic scoring, bullet rewrites, and tailoring suggestions.\\n")

    if args.jd:
        jd_text = read_text_input(args.jd, is_path=not args.jd_text)
        result  = run_jd(resume_text, jd_text, docx_issues, args.model, api_key, use_ai)
        print(json.dumps(result, indent=2)) if args.json else print_jd_report(result)
    else:
        result = run_general(resume_text, docx_issues, args.model, api_key, use_ai)
        print(json.dumps(result, indent=2)) if args.json else print_general_report(result)


if __name__ == "__main__":
    main()
